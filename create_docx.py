from docx import Document
import os

document = Document()

# Add a title
document.add_heading('1. Introducción', level=1)
document.add_paragraph(
    'El análisis de la capacidad de carga es un paso fundamental y crítico en el diseño geotécnico '
    'de cimentaciones superficiales. Su objetivo principal es garantizar que el esfuerzo transmitido por '
    'la estructura al terreno no supere la resistencia al corte del suelo, evitando así la falla por corte '
    'general o local, y asegurando que los asentamientos se mantengan dentro de rangos tolerables. '
    'En la práctica profesional moderna, el uso de herramientas computacionales para la automatización de '
    'estos cálculos resulta indispensable, ya que permite evaluar múltiples escenarios paramétricos, iterar '
    'dimensiones, analizar estratigrafías complejas y considerar condiciones variables como fluctuaciones del '
    'nivel freático de manera eficiente y precisa. El alcance de este trabajo comprende el desarrollo y '
    'validación de un aplicativo de cálculo geotécnico integral, capaz de aplicar rigurosamente las metodologías '
    'tradicionales, y el presente informe detalla el marco teórico implementado, la arquitectura de la '
    'herramienta y la validación de sus resultados.'
)

document.add_heading('2. Marco Teórico', level=1)

document.add_heading('2.1. Método de Terzaghi', level=2)
document.add_paragraph(
    'El método de Terzaghi (1943) fue la primera formulación teórica exhaustiva para evaluar la capacidad de '
    'carga en cimentaciones superficiales (Df ≤ B). Terzaghi asumió que la resistencia del suelo que se '
    'encuentra por encima del fondo de la cimentación puede ser despreciada, actuando únicamente como una '
    'sobrecarga efectiva (q). Para cimentaciones cuadradas y rectangulares (adaptadas de la condición de franja), '
    'la ecuación fundamental se expresa como:'
)
p = document.add_paragraph()
r = p.add_run("qu = 1.3 · c' · Nc + q · Nq + 0.4 · γ · B · Nγ")
r.bold = True
document.add_paragraph(
    'Donde:\n'
    '• qu: Capacidad de carga última.\n'
    '• c\': Cohesión efectiva del suelo.\n'
    '• q: Sobrecarga efectiva al nivel de desplante (q = γ · Df).\n'
    '• γ: Peso específico del suelo bajo la cimentación.\n'
    '• B: Ancho de la cimentación.\n'
    '• Nc, Nq, Nγ: Factores de capacidad de carga de Terzaghi, que dependen exclusivamente del ángulo de fricción interna (φ\').'
)

document.add_heading('2.2. Ecuación General de Capacidad de Carga', level=2)
document.add_paragraph(
    'Para superar las limitaciones del método de Terzaghi (que no considera explícitamente cimentaciones '
    'rectangulares complejas, inclinación de la carga, ni profundidad relativa de desplante), se emplea la '
    'Ecuación General de Capacidad de Carga propuesta inicialmente por Meyerhof y refinada posteriormente '
    'por Hansen, Vesic y Das. Esta formulación incorpora factores correctores empíricos y teóricos a cada uno de '
    'los términos de cohesión, sobrecarga y peso propio:'
)
p = document.add_paragraph()
r = p.add_run("qu = c' · Nc · Fcs · Fcd · Fci + q · Nq · Fqs · Fqd · Fqi + 0.5 · γ · B · Nγ · Fγs · Fγd · Fγi")
r.bold = True
document.add_paragraph(
    'Los factores de capacidad de carga para la ecuación general se definen rigurosamente como:\n'
    '• Nq = tan²(45° + φ\'/2) · e^(π tan φ\')\n'
    '• Nc = (Nq - 1) cot φ\'\n'
    '• Nγ = 2(Nq + 1) tan φ\' (Según formulación de Vesic/Das)'
)

document.add_heading('2.3. Factores Correctores', level=2)
document.add_paragraph(
    'Para la evaluación de la Ecuación General, se han implementado las siguientes relaciones empíricas (basadas en De Beer y Hansen):'
)
document.add_heading('2.3.1. Factores de Forma (Fcs, Fqs, Fγs)', level=3)
document.add_paragraph(
    'Toman en cuenta la geometría tridimensional de la falla para zapatas rectangulares (donde L > B):\n'
    '• Fcs = 1 + (B / L)(Nq / Nc)\n'
    '• Fqs = 1 + (B / L) tan φ\'\n'
    '• Fγs = 1 - 0.4(B / L)'
)

document.add_heading('2.3.2. Factores de Profundidad (Fcd, Fqd, Fγd)', level=3)
document.add_paragraph(
    'Consideran la resistencia al corte del suelo situado por encima del nivel de cimentación. Dependen de la relación Df / B:\n\n'
    'Si Df / B ≤ 1:\n'
    '• Fqd = 1 + 2 tan φ\' (1 - sin φ\')² (Df / B)\n'
    '• Fcd = Fqd - (1 - Fqd) / (Nc tan φ\')\n'
    '• Fγd = 1\n\n'
    'Si Df / B > 1:\n'
    'Se reemplaza el término (Df / B) por el arcotangente arctan(Df / B) en radianes.'
)

document.add_heading('2.3.3. Factores de Inclinación de Carga (Fci, Fqi, Fγi)', level=3)
document.add_paragraph(
    'Ajustan la capacidad de carga si la resultante de las fuerzas presenta un ángulo β con respecto a la vertical:\n'
    '• Fci = Fqi = (1 - β / 90°)²\n'
    '• Fγi = (1 - β / φ\')²\n\n'
    'Para cargas estrictamente verticales (β = 0°), todos los factores de inclinación toman un valor igual a 1.0.'
)

document.add_heading('2.4. Corrección por Nivel Freático', level=2)
document.add_paragraph(
    'La presencia de agua reduce el esfuerzo efectivo en el suelo, impactando negativamente la capacidad de soporte. '
    'La aplicación corrige el parámetro γ basándose en la posición del Nivel Freático (Dw) relativa al desplante (Df):\n\n'
    '• Caso 1 (N.F. por encima de la cimentación, 0 ≤ Dw ≤ Df):\n'
    '  El término q (sobrecarga) sufre una pérdida por subpresión. Se calcula sumando el peso efectivo (γsat - γw) '
    'de las capas sumergidas. El parámetro γ en el tercer término de la ecuación utiliza estrictamente el peso específico '
    'efectivo γ\' = γsat - γw.\n\n'
    '• Caso 2 (N.F. por debajo de la cimentación, Df < Dw ≤ Df + B):\n'
    '  La cuña de falla interseca parcialmente el nivel freático. La sobrecarga q permanece intacta con pesos secos/húmedos, '
    'pero el parámetro γ del tercer término se interpola según la profundidad d = Dw - Df:\n'
    '  γeff = γ\' + (d / B)(γ - γ\')\n\n'
    '• Caso 3 (N.F. muy profundo, Dw > Df + B):\n'
    '  El agua no influye en la zona de falla. Se asume la condición seca o de humedad natural y no se aplica ninguna reducción al peso unitario (γeff = γ).'
)

document.add_heading('3. Descripción de la Aplicación Desarrollada', level=1)
document.add_heading('3.1. Herramienta Utilizada', level=2)
document.add_paragraph(
    'La herramienta computacional "Cimentaciones Web" fue desarrollada utilizando una arquitectura moderna orientada a la web.\n'
    '• Frontend (Interfaz de Usuario): Desarrollado en React.js y TypeScript, permitiendo una experiencia de usuario interactiva '
    'y fluida. Se utilizaron librerías gráficas como Plotly.js para las iteraciones paramétricas y representaciones visuales.\n'
    '• Backend (Motor de Cálculo): Programado íntegramente en Python mediante el framework FastAPI. Se utilizaron las librerías NumPy '
    'y Math para garantizar exactitud matemática. Adicionalmente, cuenta con un motor generador de informes dinámicos y automatizados que '
    'compila directamente a código LaTeX puro, exportando reportes formales en PDF de alta calidad tipográfica.'
)

document.add_heading('3.2. Lógica y Funcionamiento', level=2)
document.add_paragraph(
    'El flujo de procesamiento del sistema opera secuencialmente de la siguiente manera:\n'
    '1. Lectura e Ingesta de Datos: El motor recibe un vector JSON con la estratigrafía del suelo, las dimensiones de la zapata, '
    'y la ubicación de las cotas de agua y sótano si aplica.\n'
    '2. Ubicación del Estrato de Diseño: El algoritmo suma iterativamente los espesores de los estratos para encontrar matemáticamente la capa '
    'donde se asienta la cota Df. Extrae automáticamente los parámetros de φ\' y c\' de esta capa.\n'
    '3. Análisis de Presiones Efectivas: Se evalúan las cotas relativas del nivel freático para computar las presiones efectivas acumuladas '
    'y determinar el peso unitario ponderado en la cuña de corte.\n'
    '4. Ejecución Matemática: Se calculan los factores de capacidad portante, forma, profundidad e inclinación en función de la metodología '
    'elegida por el usuario (Terzaghi, General, o Norma E.050).\n'
    '5. Cálculo Iterativo: En el módulo de optimización, el motor encapsula el procedimiento anterior en un bucle cerrado, iterando rangos '
    'para el ancho B y la profundidad Df con un paso Δ, retornando matrices de resultados bidimensionales que se envían al frontend para su graficación paramétrica.'
)

document.add_heading('3.3. Variables de Entrada y Salida', level=2)
document.add_paragraph(
    'Variables de Entrada (Inputs):\n'
    '• Suelo (por estrato): Espesor (H), Peso específico natural (γ), Peso específico saturado (γsat), Cohesión efectiva (c\'), Ángulo de fricción interna (φ\').\n'
    '• Cimentación: Ancho (B), Longitud (L), Profundidad de desplante (Df), Ángulo de inclinación de la carga (β), Factor de Seguridad exigido (FS).\n'
    '• Entorno: Existencia y profundidad del nivel freático (Dw), Profundidad de sótano (Ds).\n\n'
    'Variables de Salida (Outputs):\n'
    '• Intermedias: Factores Nc, Nq, Nγ; Factores geométricos y correcciones del N.F.\n'
    '• Capacidad Portante: Capacidad última (qu en tnf/m² o kPa).\n'
    '• Capacidad Admisible: Capacidad de trabajo de la cimentación (qadm = qu / FS).\n'
    '• Carga Máxima de Servicio: Carga bruta máxima permisible para la zapata (Qmax en toneladas o kN).\n'
    '• Entregables: Gráficos paramétricos iterativos, tablas resumen para MS Word/Excel, vistas del perfil de suelo en 2D y 3D, y memoria de cálculo detallada en PDF.'
)

document.add_heading('3.4. Interfaz', level=2)
document.add_paragraph(
    'La interfaz de usuario adopta un estilo profesional (tipo software CAD/Ingeniería) dividiendo el espacio de trabajo en áreas funcionales claramente definidas. '
    'En la barra lateral izquierda se disponen los "Inputs" paramétricos (dimensiones, métodos, estratigrafía) configurables mediante tablas reactivas. En el panel '
    'principal central y derecho se despliegan los "Resultados", presentando:\n'
    '1. Vistas Gráficas: Un renderizado de la sección transversal del suelo y zapatas, así como un modelo paramétrico 3D.\n'
    '2. Cuadros de Resultados Finales: Paneles destacados con los valores numéricos de qadm y Qmax.\n'
    '3. Panel de Gráficos e Iteraciones: Donde el usuario puede visualizar en tiempo real cómo cambia la resistencia del suelo variando dimensiones clave en matrices 2D interactuables, '
    'incluyendo una tabla de resumen lista para ser exportada o copiada.'
)

document.save('Informe_Cimentaciones.docx')
print('Docx generated successfully!')
